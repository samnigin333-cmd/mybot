from telegram import Update, InlineKeyboardButton, InlineKeyboardMarkup
from telegram.ext import (ContextTypes, ConversationHandler, CommandHandler,
                           MessageHandler, filters, CallbackQueryHandler)
import database as db
from datetime import datetime
import os

ADMIN_IDS = []

# States
(ADD_NAME, ADD_USERNAME, ADD_PASSWORD, ADD_POSITION,
 DEL_WORKER, KPI_SELECT, KPI_SCORE, KPI_COMMENT,
 HOURS_SELECT, HOURS_VALUE, HOURS_NOTE,
 CHANGE_PASS_SELECT, CHANGE_PASS_VALUE) = range(13)

MONTHS_UZ = {
    "01": "Yanvar", "02": "Fevral", "03": "Mart", "04": "Aprel",
    "05": "May", "06": "Iyun", "07": "Iyul", "08": "Avgust",
    "09": "Sentabr", "10": "Oktabr", "11": "Noyabr", "12": "Dekabr"
}

def is_admin(update: Update):
    return update.effective_user.id in ADMIN_IDS

def get_month_name(month_str):
    parts = month_str.split("-")
    if len(parts) == 2:
        return f"{MONTHS_UZ.get(parts[1], parts[1])} {parts[0]}"
    return month_str

def main_keyboard():
    return InlineKeyboardMarkup([
        [InlineKeyboardButton("➕ Ishchi qo'shish", callback_data="add_worker"),
         InlineKeyboardButton("👥 Ro'yxat", callback_data="list_workers")],
        [InlineKeyboardButton("❌ Ishchi o'chirish", callback_data="del_worker"),
         InlineKeyboardButton("🔑 Parol o'zgartirish", callback_data="change_pass")],
        [InlineKeyboardButton("⭐ KPI belgilash", callback_data="set_kpi"),
         InlineKeyboardButton("🕐 Ish soat belgilash", callback_data="set_hours")],
        [InlineKeyboardButton("📋 Bugungi hisobotlar", callback_data="today_reports")],
        [InlineKeyboardButton("📊 Oylik hisobot", callback_data="monthly_report"),
         InlineKeyboardButton("📅 Davomat", callback_data="attendance")],
        [InlineKeyboardButton("💾 Eksport", callback_data="export_menu")],
    ])

async def admin_menu(update: Update, context: ContextTypes.DEFAULT_TYPE):
    if not is_admin(update):
        await update.message.reply_text("🚫 Sizda ruxsat yo'q.")
        return
    now = datetime.now().strftime("%d.%m.%Y %H:%M")
    workers = db.get_all_workers()
    await update.message.reply_text(
        f"👨‍💼 *ADMIN PANEL*\n"
        f"─────────────────\n"
        f"📅 {now}\n"
        f"👥 Jami ishchilar: *{len(workers)} ta*\n"
        f"─────────────────\n"
        f"Quyidagi amallardan birini tanlang:",
        parse_mode="Markdown",
        reply_markup=main_keyboard()
    )

async def admin_menu_callback(update: Update, context: ContextTypes.DEFAULT_TYPE):
    query = update.callback_query
    await query.answer()
    now = datetime.now().strftime("%d.%m.%Y %H:%M")
    workers = db.get_all_workers()
    await query.edit_message_text(
        f"👨‍💼 *ADMIN PANEL*\n"
        f"─────────────────\n"
        f"📅 {now}\n"
        f"👥 Jami ishchilar: *{len(workers)} ta*\n"
        f"─────────────────\n"
        f"Quyidagi amallardan birini tanlang:",
        parse_mode="Markdown",
        reply_markup=main_keyboard()
    )

async def admin_callback(update: Update, context: ContextTypes.DEFAULT_TYPE):
    query = update.callback_query
    await query.answer()
    data = query.data

    if data == "list_workers":
        workers = db.get_all_workers()
        if not workers:
            await query.edit_message_text(
                "📭 Hozircha ishchilar yo'q.\n\n"
                "Ishchi qo'shish uchun /admin → ➕ Ishchi qo'shish",
                reply_markup=InlineKeyboardMarkup([[InlineKeyboardButton("🔙 Orqaga", callback_data="admin_back")]])
            )
            return
        text = f"👥 *ISHCHILAR RO'YXATI* ({len(workers)} ta)\n{'─'*25}\n\n"
        for i, w in enumerate(workers, 1):
            linked = "✅ ulangan" if w['telegram_id'] else "⏳ ulanmagan"
            text += (f"*{i}. {w['full_name']}*\n"
                     f"   🔑 Login: `{w['username']}`\n"
                     f"   💼 Lavozim: {w['position'] or '—'}\n"
                     f"   📱 Telegram: {linked}\n\n")
        await query.edit_message_text(
            text, parse_mode="Markdown",
            reply_markup=InlineKeyboardMarkup([[InlineKeyboardButton("🔙 Orqaga", callback_data="admin_back")]])
        )

    elif data == "today_reports":
        reports = db.get_reports_today()
        today_str = datetime.now().strftime("%d.%m.%Y")
        if not reports:
            await query.edit_message_text(
                f"📭 {today_str} kuni hali hisobot yo'q.",
                reply_markup=InlineKeyboardMarkup([[InlineKeyboardButton("🔙 Orqaga", callback_data="admin_back")]])
            )
            return
        text = f"📋 *BUGUNGI HISOBOTLAR* ({today_str})\n{'─'*25}\n\n"
        for r in reports:
            time_str = r['created_at'][11:16] if r['created_at'] else ""
            photo_mark = " 📸" if r['photo_file_id'] else ""
            text += f"👤 *{r['full_name']}*{photo_mark} — {time_str}\n{r['report_text']}\n\n"
        await query.edit_message_text(
            text, parse_mode="Markdown",
            reply_markup=InlineKeyboardMarkup([[InlineKeyboardButton("🔙 Orqaga", callback_data="admin_back")]])
        )

    elif data == "monthly_report":
        month = datetime.now().strftime("%Y-%m")
        month_name = get_month_name(month)
        reports = db.get_reports_by_month(month)
        kpis = db.get_kpi_by_month(month)
        hours = db.get_work_hours_month(month)
        workers = db.get_all_workers()

        text = f"📊 *{month_name.upper()} OYLIK HISOBOT*\n{'─'*25}\n\n"
        text += f"👥 Ishchilar: *{len(workers)} ta*\n"
        text += f"📝 Hisobotlar: *{len(reports)} ta*\n"
        text += f"⭐ KPI yozuvlar: *{len(kpis)} ta*\n\n"

        if kpis:
            text += "⭐ *KPI BALLARI:*\n"
            for k in kpis:
                text += f"• {k['full_name']}: *{k['score']}/10*"
                if k['comment']:
                    text += f" — {k['comment']}"
                text += "\n"
            text += "\n"

        if hours:
            text += "🕐 *ISH SOATLARI:*\n"
            for h in hours:
                text += f"• {h['full_name']}: *{h['hours']} soat*"
                if h['note']:
                    text += f" ({h['note']})"
                text += "\n"

        await query.edit_message_text(
            text, parse_mode="Markdown",
            reply_markup=InlineKeyboardMarkup([[InlineKeyboardButton("🔙 Orqaga", callback_data="admin_back")]])
        )

    elif data == "attendance":
        month = datetime.now().strftime("%Y-%m")
        month_name = get_month_name(month)
        records = db.get_attendance_month(month)
        if not records:
            await query.edit_message_text(
                f"📭 {month_name} oyida davomat ma'lumoti yo'q.",
                reply_markup=InlineKeyboardMarkup([[InlineKeyboardButton("🔙 Orqaga", callback_data="admin_back")]])
            )
            return
        text = f"📅 *DAVOMAT — {month_name.upper()}*\n{'─'*25}\n\n"
        for r in records:
            emoji = "🟢" if r['att_type'] == "in" else "🔴"
            dt = r['att_datetime'][:16] if r['att_datetime'] else ""
            text += f"{emoji} *{r['full_name']}* — {dt}\n"
            if r['note']:
                text += f"   📝 {r['note']}\n"
        await query.edit_message_text(
            text, parse_mode="Markdown",
            reply_markup=InlineKeyboardMarkup([[InlineKeyboardButton("🔙 Orqaga", callback_data="admin_back")]])
        )

    elif data == "export_menu":
        await query.edit_message_text(
            "💾 *EKSPORT*\n\nQaysi formatda saqlash kerak?",
            parse_mode="Markdown",
            reply_markup=InlineKeyboardMarkup([
                [InlineKeyboardButton("📄 Word (.docx)", callback_data="export_word")],
                [InlineKeyboardButton("📊 Excel (.xlsx)", callback_data="export_excel")],
                [InlineKeyboardButton("📑 PDF (.pdf)", callback_data="export_pdf")],
                [InlineKeyboardButton("🔙 Orqaga", callback_data="admin_back")],
            ])
        )

    elif data == "export_word":
        await query.edit_message_text("⏳ Word fayl tayyorlanmoqda...")
        try:
            filepath = await generate_word_report()
            await query.message.reply_document(
                document=open(filepath, 'rb'),
                filename=os.path.basename(filepath),
                caption=f"📄 Oylik hisobot — {get_month_name(datetime.now().strftime('%Y-%m'))}"
            )
            await query.edit_message_text(
                "✅ Word fayl yuborildi!",
                reply_markup=InlineKeyboardMarkup([[InlineKeyboardButton("🔙 Orqaga", callback_data="admin_back")]])
            )
        except Exception as e:
            await query.edit_message_text(
                f"❌ Xato: {str(e)}",
                reply_markup=InlineKeyboardMarkup([[InlineKeyboardButton("🔙 Orqaga", callback_data="admin_back")]])
            )

    elif data == "export_excel":
        await query.edit_message_text("⏳ Excel fayl tayyorlanmoqda...")
        try:
            filepath = await generate_excel_report()
            await query.message.reply_document(
                document=open(filepath, 'rb'),
                filename=os.path.basename(filepath),
                caption=f"📊 Oylik hisobot — {get_month_name(datetime.now().strftime('%Y-%m'))}"
            )
            await query.edit_message_text(
                "✅ Excel fayl yuborildi!",
                reply_markup=InlineKeyboardMarkup([[InlineKeyboardButton("🔙 Orqaga", callback_data="admin_back")]])
            )
        except Exception as e:
            await query.edit_message_text(
                f"❌ Xato: {str(e)}",
                reply_markup=InlineKeyboardMarkup([[InlineKeyboardButton("🔙 Orqaga", callback_data="admin_back")]])
            )

    elif data == "export_pdf":
        await query.edit_message_text("⏳ PDF fayl tayyorlanmoqda...")
        try:
            filepath = await generate_pdf_report()
            await query.message.reply_document(
                document=open(filepath, 'rb'),
                filename=os.path.basename(filepath),
                caption=f"📑 Oylik hisobot — {get_month_name(datetime.now().strftime('%Y-%m'))}"
            )
            await query.edit_message_text(
                "✅ PDF fayl yuborildi!",
                reply_markup=InlineKeyboardMarkup([[InlineKeyboardButton("🔙 Orqaga", callback_data="admin_back")]])
            )
        except Exception as e:
            await query.edit_message_text(
                f"❌ Xato: {str(e)}",
                reply_markup=InlineKeyboardMarkup([[InlineKeyboardButton("🔙 Orqaga", callback_data="admin_back")]])
            )

    elif data == "admin_back":
        await admin_menu_callback(update, context)

# ==================== EXPORT FUNCTIONS ====================

async def generate_excel_report():
    import openpyxl
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side

    month = datetime.now().strftime("%Y-%m")
    month_name = get_month_name(month)
    workers = db.get_all_workers()
    kpis = db.get_kpi_by_month(month)
    hours = db.get_work_hours_month(month)
    reports = db.get_reports_by_month(month)
    attendance = db.get_attendance_month(month)

    wb = openpyxl.Workbook()

    # Style helpers
    header_font = Font(bold=True, color="FFFFFF", size=12)
    header_fill = PatternFill("solid", fgColor="1F4E79")
    center = Alignment(horizontal="center", vertical="center", wrap_text=True)
    left = Alignment(horizontal="left", vertical="center", wrap_text=True)
    thin = Border(
        left=Side(style='thin'), right=Side(style='thin'),
        top=Side(style='thin'), bottom=Side(style='thin')
    )

    def style_header(cell, fill_color="1F4E79"):
        cell.font = Font(bold=True, color="FFFFFF", size=11)
        cell.fill = PatternFill("solid", fgColor=fill_color)
        cell.alignment = center
        cell.border = thin

    def style_cell(cell, align=None):
        cell.border = thin
        cell.alignment = align or left

    # ===== Sheet 1: Umumiy =====
    ws = wb.active
    ws.title = "Umumiy"
    ws.merge_cells("A1:E1")
    ws["A1"] = f"OYLIK HISOBOT — {month_name.upper()}"
    ws["A1"].font = Font(bold=True, size=14, color="1F4E79")
    ws["A1"].alignment = center

    ws.merge_cells("A2:E2")
    ws["A2"] = f"Tuzilgan: {datetime.now().strftime('%d.%m.%Y %H:%M')}"
    ws["A2"].alignment = center

    headers = ["#", "Ism Familiya", "Lavozim", "Ish soat", "KPI ball"]
    for col, h in enumerate(headers, 1):
        style_header(ws.cell(row=4, column=col))
        ws.cell(row=4, column=col).value = h

    kpi_dict = {k['full_name']: k['score'] for k in kpis}
    hours_dict = {h['full_name']: h['hours'] for h in hours}

    for i, w in enumerate(workers, 1):
        row = 4 + i
        data = [i, w['full_name'], w['position'] or "—",
                hours_dict.get(w['full_name'], "—"),
                kpi_dict.get(w['full_name'], "—")]
        for col, val in enumerate(data, 1):
            c = ws.cell(row=row, column=col, value=val)
            style_cell(c, center if col in [1, 4, 5] else left)
            if i % 2 == 0:
                c.fill = PatternFill("solid", fgColor="EBF3FB")

    ws.column_dimensions['A'].width = 5
    ws.column_dimensions['B'].width = 25
    ws.column_dimensions['C'].width = 20
    ws.column_dimensions['D'].width = 12
    ws.column_dimensions['E'].width = 12

    # ===== Sheet 2: KPI =====
    ws2 = wb.create_sheet("KPI")
    ws2.merge_cells("A1:D1")
    ws2["A1"] = f"KPI BALLARI — {month_name.upper()}"
    ws2["A1"].font = Font(bold=True, size=13, color="1F4E79")
    ws2["A1"].alignment = center

    for col, h in enumerate(["#", "Ism Familiya", "Ball (1-10)", "Izoh"], 1):
        style_header(ws2.cell(row=3, column=col), "1F4E79")
        ws2.cell(row=3, column=col).value = h

    for i, k in enumerate(kpis, 1):
        row = 3 + i
        for col, val in enumerate([i, k['full_name'], k['score'], k['comment']], 1):
            c = ws2.cell(row=row, column=col, value=val)
            style_cell(c, center if col in [1, 3] else left)
            if i % 2 == 0:
                c.fill = PatternFill("solid", fgColor="EBF3FB")
            if col == 3 and isinstance(val, int):
                if val <= 4:
                    c.font = Font(color="C00000", bold=True)
                elif val >= 8:
                    c.font = Font(color="375623", bold=True)

    ws2.column_dimensions['A'].width = 5
    ws2.column_dimensions['B'].width = 25
    ws2.column_dimensions['C'].width = 12
    ws2.column_dimensions['D'].width = 40

    # ===== Sheet 3: Hisobotlar =====
    ws3 = wb.create_sheet("Hisobotlar")
    ws3.merge_cells("A1:D1")
    ws3["A1"] = f"KUNLIK HISOBOTLAR — {month_name.upper()}"
    ws3["A1"].font = Font(bold=True, size=13, color="1F4E79")
    ws3["A1"].alignment = center

    for col, h in enumerate(["#", "Sana", "Ism Familiya", "Hisobot"], 1):
        style_header(ws3.cell(row=3, column=col), "1F4E79")
        ws3.cell(row=3, column=col).value = h

    for i, r in enumerate(reports, 1):
        row = 3 + i
        for col, val in enumerate([i, r['date'], r['full_name'], r['report_text']], 1):
            c = ws3.cell(row=row, column=col, value=val)
            style_cell(c, center if col in [1, 2] else left)
            if i % 2 == 0:
                c.fill = PatternFill("solid", fgColor="EBF3FB")

    ws3.column_dimensions['A'].width = 5
    ws3.column_dimensions['B'].width = 12
    ws3.column_dimensions['C'].width = 25
    ws3.column_dimensions['D'].width = 50
    ws3.row_dimensions[1].height = 30

    # ===== Sheet 4: Davomat =====
    ws4 = wb.create_sheet("Davomat")
    ws4.merge_cells("A1:D1")
    ws4["A1"] = f"DAVOMAT — {month_name.upper()}"
    ws4["A1"].font = Font(bold=True, size=13, color="1F4E79")
    ws4["A1"].alignment = center

    for col, h in enumerate(["#", "Ism Familiya", "Tur", "Vaqt"], 1):
        style_header(ws4.cell(row=3, column=col), "1F4E79")
        ws4.cell(row=3, column=col).value = h

    for i, a in enumerate(attendance, 1):
        row = 3 + i
        att_label = "Keldi 🟢" if a['att_type'] == "in" else "Ketdi 🔴"
        for col, val in enumerate([i, a['full_name'], att_label, a['att_datetime'][:16]], 1):
            c = ws4.cell(row=row, column=col, value=val)
            style_cell(c, center if col in [1, 3, 4] else left)

    ws4.column_dimensions['A'].width = 5
    ws4.column_dimensions['B'].width = 25
    ws4.column_dimensions['C'].width = 12
    ws4.column_dimensions['D'].width = 20

    os.makedirs("exports", exist_ok=True)
    filename = f"exports/hisobot_{month}.xlsx"
    wb.save(filename)
    return filename

async def generate_word_report():
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_ALIGN_VERTICAL
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    month = datetime.now().strftime("%Y-%m")
    month_name = get_month_name(month)
    workers = db.get_all_workers()
    kpis = db.get_kpi_by_month(month)
    hours = db.get_work_hours_month(month)
    reports = db.get_reports_by_month(month)

    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # Title
    title = doc.add_heading(f"OYLIK HISOBOT — {month_name.upper()}", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.runs[0].font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

    subtitle = doc.add_paragraph(f"Tuzilgan: {datetime.now().strftime('%d.%m.%Y %H:%M')}")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].font.size = Pt(11)
    subtitle.runs[0].font.color.rgb = RGBColor(0x70, 0x70, 0x70)

    doc.add_paragraph()

    # Summary
    doc.add_heading("UMUMIY MA'LUMOT", 1)
    summary_data = [
        ("Jami ishchilar:", str(len(workers))),
        ("Hisobotlar soni:", str(len(reports))),
        ("KPI yozuvlar:", str(len(kpis))),
        ("Oy:", month_name),
    ]
    for label, val in summary_data:
        p = doc.add_paragraph()
        run1 = p.add_run(f"{label} ")
        run1.bold = True
        p.add_run(val)

    doc.add_paragraph()

    # KPI section
    if kpis:
        doc.add_heading("KPI BALLARI", 1)
        table = doc.add_table(rows=1, cols=3)
        table.style = 'Table Grid'
        hdr = table.rows[0].cells
        for i, text in enumerate(["Ism Familiya", "Ball (1-10)", "Izoh"]):
            hdr[i].text = text
            hdr[i].paragraphs[0].runs[0].bold = True

        for k in kpis:
            row = table.add_row().cells
            row[0].text = k['full_name']
            row[1].text = str(k['score'])
            row[2].text = k['comment'] or "—"

        doc.add_paragraph()

    # Work hours section
    if hours:
        doc.add_heading("ISH SOATLARI", 1)
        table2 = doc.add_table(rows=1, cols=3)
        table2.style = 'Table Grid'
        hdr2 = table2.rows[0].cells
        for i, text in enumerate(["Ism Familiya", "Soat soni", "Izoh"]):
            hdr2[i].text = text
            hdr2[i].paragraphs[0].runs[0].bold = True

        for h in hours:
            row = table2.add_row().cells
            row[0].text = h['full_name']
            row[1].text = str(h['hours'])
            row[2].text = h['note'] or "—"

        doc.add_paragraph()

    # Reports section
    if reports:
        doc.add_heading("KUNLIK HISOBOTLAR", 1)
        for r in reports:
            p = doc.add_paragraph()
            run = p.add_run(f"📅 {r['date']} — {r['full_name']}: ")
            run.bold = True
            p.add_run(r['report_text'])

    os.makedirs("exports", exist_ok=True)
    filename = f"exports/hisobot_{month}.docx"
    doc.save(filename)
    return filename

async def generate_pdf_report():
    from reportlab.lib.pagesizes import A4
    from reportlab.lib import colors
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import cm
    from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont
    import urllib.request

    month = datetime.now().strftime("%Y-%m")
    month_name = get_month_name(month)
    workers = db.get_all_workers()
    kpis = db.get_kpi_by_month(month)
    hours = db.get_work_hours_month(month)
    reports = db.get_reports_by_month(month)

    os.makedirs("exports", exist_ok=True)
    filename = f"exports/hisobot_{month}.pdf"

    doc = SimpleDocTemplate(filename, pagesize=A4,
                             topMargin=2*cm, bottomMargin=2*cm,
                             leftMargin=2.5*cm, rightMargin=2.5*cm)

    styles = getSampleStyleSheet()
    title_style = ParagraphStyle('title', parent=styles['Title'],
                                  fontSize=16, textColor=colors.HexColor('#1F4E79'))
    h1_style = ParagraphStyle('h1', parent=styles['Heading1'],
                               fontSize=13, textColor=colors.HexColor('#1F4E79'))
    normal = styles['Normal']

    content = []

    # Title
    content.append(Paragraph(f"OYLIK HISOBOT — {month_name.upper()}", title_style))
    content.append(Paragraph(f"Tuzilgan: {datetime.now().strftime('%d.%m.%Y %H:%M')}", normal))
    content.append(Spacer(1, 0.5*cm))

    # Summary
    content.append(Paragraph("UMUMIY MA'LUMOT", h1_style))
    content.append(Paragraph(f"Jami ishchilar: {len(workers)}", normal))
    content.append(Paragraph(f"Hisobotlar: {len(reports)}", normal))
    content.append(Paragraph(f"KPI yozuvlar: {len(kpis)}", normal))
    content.append(Spacer(1, 0.5*cm))

    # KPI table
    if kpis:
        content.append(Paragraph("KPI BALLARI", h1_style))
        kpi_data = [["Ism Familiya", "Ball", "Izoh"]]
        for k in kpis:
            kpi_data.append([k['full_name'], str(k['score']), k['comment'] or "—"])
        kpi_table = Table(kpi_data, colWidths=[6*cm, 2.5*cm, 8*cm])
        kpi_table.setStyle(TableStyle([
            ('BACKGROUND', (0,0), (-1,0), colors.HexColor('#1F4E79')),
            ('TEXTCOLOR', (0,0), (-1,0), colors.white),
            ('FONTNAME', (0,0), (-1,0), 'Helvetica-Bold'),
            ('FONTSIZE', (0,0), (-1,-1), 10),
            ('GRID', (0,0), (-1,-1), 0.5, colors.grey),
            ('ROWBACKGROUNDS', (0,1), (-1,-1), [colors.white, colors.HexColor('#EBF3FB')]),
            ('ALIGN', (1,0), (1,-1), 'CENTER'),
            ('VALIGN', (0,0), (-1,-1), 'MIDDLE'),
            ('TOPPADDING', (0,0), (-1,-1), 6),
            ('BOTTOMPADDING', (0,0), (-1,-1), 6),
        ]))
        content.append(kpi_table)
        content.append(Spacer(1, 0.5*cm))

    # Work hours table
    if hours:
        content.append(Paragraph("ISH SOATLARI", h1_style))
        h_data = [["Ism Familiya", "Soat", "Izoh"]]
        for h in hours:
            h_data.append([h['full_name'], str(h['hours']), h['note'] or "—"])
        h_table = Table(h_data, colWidths=[6*cm, 2.5*cm, 8*cm])
        h_table.setStyle(TableStyle([
            ('BACKGROUND', (0,0), (-1,0), colors.HexColor('#1F4E79')),
            ('TEXTCOLOR', (0,0), (-1,0), colors.white),
            ('FONTNAME', (0,0), (-1,0), 'Helvetica-Bold'),
            ('FONTSIZE', (0,0), (-1,-1), 10),
            ('GRID', (0,0), (-1,-1), 0.5, colors.grey),
            ('ROWBACKGROUNDS', (0,1), (-1,-1), [colors.white, colors.HexColor('#EBF3FB')]),
            ('VALIGN', (0,0), (-1,-1), 'MIDDLE'),
            ('TOPPADDING', (0,0), (-1,-1), 6),
            ('BOTTOMPADDING', (0,0), (-1,-1), 6),
        ]))
        content.append(h_table)
        content.append(Spacer(1, 0.5*cm))

    # Reports
    if reports:
        content.append(Paragraph("KUNLIK HISOBOTLAR", h1_style))
        for r in reports:
            content.append(Paragraph(
                f"<b>{r['date']} — {r['full_name']}:</b> {r['report_text']}", normal
            ))

    doc.build(content)
    return filename

# ==================== CONVERSATION HANDLERS ====================

async def add_worker_start(update: Update, context: ContextTypes.DEFAULT_TYPE):
    query = update.callback_query
    await query.answer()
    await query.edit_message_text(
        "➕ *YANGI ISHCHI QO'SHISH*\n\n"
        "1️⃣ Ishchining to'liq ism-sharifini kiriting:\n\n"
        "_Bekor qilish: /cancel_",
        parse_mode="Markdown"
    )
    return ADD_NAME

async def add_name(update: Update, context: ContextTypes.DEFAULT_TYPE):
    name = update.message.text.strip()
    if len(name) < 3:
        await update.message.reply_text("❌ Ism juda qisqa. Qayta kiriting:")
        return ADD_NAME
    context.user_data['new_name'] = name
    await update.message.reply_text(
        f"✅ Ism: *{name}*\n\n"
        "2️⃣ Login (username) kiriting:\n"
        "_Faqat lotin harflari va raqamlar_",
        parse_mode="Markdown"
    )
    return ADD_USERNAME

async def add_username(update: Update, context: ContextTypes.DEFAULT_TYPE):
    username = update.message.text.strip().lower()
    if len(username) < 3:
        await update.message.reply_text("❌ Login juda qisqa (min 3 belgi). Qayta kiriting:")
        return ADD_USERNAME
    context.user_data['new_username'] = username
    await update.message.reply_text(
        f"✅ Login: `{username}`\n\n"
        "3️⃣ Parol kiriting:\n"
        "_Kamida 4 belgi_",
        parse_mode="Markdown"
    )
    return ADD_PASSWORD

async def add_password(update: Update, context: ContextTypes.DEFAULT_TYPE):
    password = update.message.text.strip()
    if len(password) < 4:
        await update.message.reply_text("❌ Parol juda qisqa (min 4 belgi). Qayta kiriting:")
        return ADD_PASSWORD
    context.user_data['new_password'] = password
    await update.message.reply_text(
        "4️⃣ Lavozimni kiriting:\n"
        "_Ixtiyoriy. O'tkazib yuborish uchun — yuboring_",
        parse_mode="Markdown"
    )
    return ADD_POSITION

async def add_position(update: Update, context: ContextTypes.DEFAULT_TYPE):
    position = update.message.text.strip()
    if position == "—":
        position = ""
    success = db.add_worker(
        context.user_data['new_name'],
        context.user_data['new_username'],
        context.user_data['new_password'],
        position
    )
    if success:
        await update.message.reply_text(
            f"✅ *Ishchi muvaffaqiyatli qo'shildi!*\n\n"
            f"👤 Ism: *{context.user_data['new_name']}*\n"
            f"🔑 Login: `{context.user_data['new_username']}`\n"
            f"🔐 Parol: `{context.user_data['new_password']}`\n"
            f"💼 Lavozim: {position or '—'}\n\n"
            f"_Ishchi /login buyrug'i bilan tizimga kirishi mumkin._",
            parse_mode="Markdown",
            reply_markup=InlineKeyboardMarkup([[InlineKeyboardButton("🔙 Admin panel", callback_data="admin_back")]])
        )
    else:
        await update.message.reply_text(
            f"❌ `{context.user_data['new_username']}` login band. Boshqa login bilan qayta urinib ko'ring.\n\n/admin",
            parse_mode="Markdown"
        )
    context.user_data.clear()
    return ConversationHandler.END

async def del_worker_start(update: Update, context: ContextTypes.DEFAULT_TYPE):
    query = update.callback_query
    await query.answer()
    workers = db.get_all_workers()
    if not workers:
        await query.edit_message_text("📭 Ishchilar yo'q.")
        return ConversationHandler.END
    keyboard = [[InlineKeyboardButton(f"❌ {w['full_name']}", callback_data=f"del_{w['id']}")] for w in workers]
    keyboard.append([InlineKeyboardButton("🔙 Orqaga", callback_data="admin_back")])
    await query.edit_message_text(
        "🗑 *ISHCHI O'CHIRISH*\n\nQaysi ishchini o'chirmoqchisiz?",
        parse_mode="Markdown",
        reply_markup=InlineKeyboardMarkup(keyboard)
    )
    return DEL_WORKER

async def del_worker_confirm(update: Update, context: ContextTypes.DEFAULT_TYPE):
    query = update.callback_query
    await query.answer()
    worker_id = int(query.data.split("_")[1])
    worker = db.get_worker_by_id(worker_id)
    if worker:
        db.delete_worker(worker_id)
        await query.edit_message_text(
            f"✅ *{worker['full_name']}* tizimdan o'chirildi.",
            parse_mode="Markdown",
            reply_markup=InlineKeyboardMarkup([[InlineKeyboardButton("🔙 Admin panel", callback_data="admin_back")]])
        )
    return ConversationHandler.END

async def kpi_start(update: Update, context: ContextTypes.DEFAULT_TYPE):
    query = update.callback_query
    await query.answer()
    workers = db.get_all_workers()
    if not workers:
        await query.edit_message_text("📭 Ishchilar yo'q.")
        return ConversationHandler.END
    keyboard = [[InlineKeyboardButton(f"👤 {w['full_name']}", callback_data=f"kpi_{w['id']}")] for w in workers]
    keyboard.append([InlineKeyboardButton("🔙 Orqaga", callback_data="admin_back")])
    await query.edit_message_text(
        "⭐ *KPI BELGILASH*\n\nQaysi ishchiga KPI qo'yasiz?",
        parse_mode="Markdown",
        reply_markup=InlineKeyboardMarkup(keyboard)
    )
    return KPI_SELECT

async def kpi_select(update: Update, context: ContextTypes.DEFAULT_TYPE):
    query = update.callback_query
    await query.answer()
    worker_id = int(query.data.split("_")[1])
    worker = db.get_worker_by_id(worker_id)
    context.user_data['kpi_worker_id'] = worker_id
    context.user_data['kpi_worker_name'] = worker['full_name']

    # Oldingi KPI ko'rsatish
    current_kpis = db.get_worker_kpi(worker_id)
    prev_text = ""
    if current_kpis:
        last = current_kpis[-1]
        prev_text = f"\n📊 Oldingi ball: *{last['score']}/10* — {last['comment']}"

    await query.edit_message_text(
        f"⭐ *{worker['full_name']}* uchun KPI{prev_text}\n\n"
        f"Ball kiriting *(1-10)*:",
        parse_mode="Markdown"
    )
    return KPI_SCORE

async def kpi_score(update: Update, context: ContextTypes.DEFAULT_TYPE):
    try:
        score = int(update.message.text.strip())
        if not 1 <= score <= 10:
            raise ValueError
        context.user_data['kpi_score'] = score
        emoji = "🔴" if score <= 4 else ("🟡" if score <= 6 else "🟢")
        await update.message.reply_text(
            f"{emoji} Ball: *{score}/10*\n\n"
            f"📝 Izoh yozing:\n"
            f"_(KPI kamaytirilsa sabab yozing)_",
            parse_mode="Markdown"
        )
        return KPI_COMMENT
    except ValueError:
        await update.message.reply_text("❌ 1 dan 10 gacha son kiriting:")
        return KPI_SCORE

async def kpi_comment(update: Update, context: ContextTypes.DEFAULT_TYPE):
    comment = update.message.text.strip()
    worker_id = context.user_data['kpi_worker_id']
    score = context.user_data['kpi_score']
    name = context.user_data['kpi_worker_name']

    db.add_kpi(worker_id, score, comment, set_by=update.effective_user.id)

    emoji = "🔴" if score <= 4 else ("🟡" if score <= 6 else "🟢")
    await update.message.reply_text(
        f"✅ *KPI saqlandi!*\n\n"
        f"👤 Ishchi: *{name}*\n"
        f"{emoji} Ball: *{score}/10*\n"
        f"📝 Izoh: {comment}",
        parse_mode="Markdown",
        reply_markup=InlineKeyboardMarkup([[InlineKeyboardButton("🔙 Admin panel", callback_data="admin_back")]])
    )
    context.user_data.clear()
    return ConversationHandler.END

async def hours_start(update: Update, context: ContextTypes.DEFAULT_TYPE):
    query = update.callback_query
    await query.answer()
    workers = db.get_all_workers()
    if not workers:
        await query.edit_message_text("📭 Ishchilar yo'q.")
        return ConversationHandler.END
    keyboard = [[InlineKeyboardButton(f"👤 {w['full_name']}", callback_data=f"hours_{w['id']}")] for w in workers]
    keyboard.append([InlineKeyboardButton("🔙 Orqaga", callback_data="admin_back")])
    await query.edit_message_text(
        "🕐 *ISH SOAT BELGILASH*\n\nQaysi ishchi uchun soat belgilaysiz?",
        parse_mode="Markdown",
        reply_markup=InlineKeyboardMarkup(keyboard)
    )
    return HOURS_SELECT

async def hours_select(update: Update, context: ContextTypes.DEFAULT_TYPE):
    query = update.callback_query
    await query.answer()
    worker_id = int(query.data.split("_")[1])
    worker = db.get_worker_by_id(worker_id)
    context.user_data['hours_worker_id'] = worker_id
    context.user_data['hours_worker_name'] = worker['full_name']

    current = db.get_worker_hours(worker_id)
    prev_text = f"\n📊 Joriy oy soati: *{current['hours']}*" if current else ""

    await query.edit_message_text(
        f"🕐 *{worker['full_name']}* uchun ish soati{prev_text}\n\n"
        f"Soat sonini kiriting *(masalan: 176)*:",
        parse_mode="Markdown"
    )
    return HOURS_VALUE

async def hours_value(update: Update, context: ContextTypes.DEFAULT_TYPE):
    try:
        hours = float(update.message.text.strip().replace(",", "."))
        if hours <= 0 or hours > 744:
            raise ValueError
        context.user_data['hours_value'] = hours
        await update.message.reply_text(
            f"⏰ Soat: *{hours}*\n\n"
            f"📝 Izoh yozing _(ixtiyoriy, o'tkazish uchun — yuboring)_:",
            parse_mode="Markdown"
        )
        return HOURS_NOTE
    except ValueError:
        await update.message.reply_text("❌ To'g'ri son kiriting (1-744):")
        return HOURS_VALUE

async def hours_note(update: Update, context: ContextTypes.DEFAULT_TYPE):
    note = update.message.text.strip()
    if note == "—":
        note = ""
    worker_id = context.user_data['hours_worker_id']
    hours = context.user_data['hours_value']
    name = context.user_data['hours_worker_name']

    db.set_work_hours(worker_id, hours, note=note, set_by=update.effective_user.id)

    await update.message.reply_text(
        f"✅ *Ish soati saqlandi!*\n\n"
        f"👤 Ishchi: *{name}*\n"
        f"⏰ Soat: *{hours}*\n"
        f"📝 Izoh: {note or '—'}",
        parse_mode="Markdown",
        reply_markup=InlineKeyboardMarkup([[InlineKeyboardButton("🔙 Admin panel", callback_data="admin_back")]])
    )
    context.user_data.clear()
    return ConversationHandler.END

async def change_pass_start(update: Update, context: ContextTypes.DEFAULT_TYPE):
    query = update.callback_query
    await query.answer()
    workers = db.get_all_workers()
    if not workers:
        await query.edit_message_text("📭 Ishchilar yo'q.")
        return ConversationHandler.END
    keyboard = [[InlineKeyboardButton(f"🔑 {w['full_name']}", callback_data=f"chp_{w['id']}")] for w in workers]
    keyboard.append([InlineKeyboardButton("🔙 Orqaga", callback_data="admin_back")])
    await query.edit_message_text(
        "🔑 *PAROL O'ZGARTIRISH*\n\nQaysi ishchi parolini o'zgartirasiz?",
        parse_mode="Markdown",
        reply_markup=InlineKeyboardMarkup(keyboard)
    )
    return CHANGE_PASS_SELECT

async def change_pass_select(update: Update, context: ContextTypes.DEFAULT_TYPE):
    query = update.callback_query
    await query.answer()
    worker_id = int(query.data.split("_")[1])
    worker = db.get_worker_by_id(worker_id)
    context.user_data['chp_worker_id'] = worker_id
    context.user_data['chp_worker_name'] = worker['full_name']
    await query.edit_message_text(
        f"🔑 *{worker['full_name']}* uchun yangi parol kiriting:",
        parse_mode="Markdown"
    )
    return CHANGE_PASS_VALUE

async def change_pass_value(update: Update, context: ContextTypes.DEFAULT_TYPE):
    new_pass = update.message.text.strip()
    if len(new_pass) < 4:
        await update.message.reply_text("❌ Parol kamida 4 belgi bo'lishi kerak:")
        return CHANGE_PASS_VALUE
    worker_id = context.user_data['chp_worker_id']
    name = context.user_data['chp_worker_name']
    db.update_worker_password(worker_id, new_pass)
    await update.message.reply_text(
        f"✅ *{name}* parol o'zgartirildi!\n"
        f"🔐 Yangi parol: `{new_pass}`",
        parse_mode="Markdown",
        reply_markup=InlineKeyboardMarkup([[InlineKeyboardButton("🔙 Admin panel", callback_data="admin_back")]])
    )
    context.user_data.clear()
    return ConversationHandler.END

async def cancel(update: Update, context: ContextTypes.DEFAULT_TYPE):
    context.user_data.clear()
    await update.message.reply_text(
        "❌ Bekor qilindi.\n\n/admin — Admin panel",
        reply_markup=InlineKeyboardMarkup([[InlineKeyboardButton("🔙 Admin panel", callback_data="admin_back")]])
    )
    return ConversationHandler.END

def get_admin_conv_handler():
    return ConversationHandler(
        entry_points=[
            CallbackQueryHandler(add_worker_start, pattern="^add_worker$"),
            CallbackQueryHandler(del_worker_start, pattern="^del_worker$"),
            CallbackQueryHandler(kpi_start, pattern="^set_kpi$"),
            CallbackQueryHandler(hours_start, pattern="^set_hours$"),
            CallbackQueryHandler(change_pass_start, pattern="^change_pass$"),
        ],
        states={
            ADD_NAME: [MessageHandler(filters.TEXT & ~filters.COMMAND, add_name)],
            ADD_USERNAME: [MessageHandler(filters.TEXT & ~filters.COMMAND, add_username)],
            ADD_PASSWORD: [MessageHandler(filters.TEXT & ~filters.COMMAND, add_password)],
            ADD_POSITION: [MessageHandler(filters.TEXT & ~filters.COMMAND, add_position)],
            DEL_WORKER: [CallbackQueryHandler(del_worker_confirm, pattern="^del_\\d+$")],
            KPI_SELECT: [CallbackQueryHandler(kpi_select, pattern="^kpi_\\d+$")],
            KPI_SCORE: [MessageHandler(filters.TEXT & ~filters.COMMAND, kpi_score)],
            KPI_COMMENT: [MessageHandler(filters.TEXT & ~filters.COMMAND, kpi_comment)],
            HOURS_SELECT: [CallbackQueryHandler(hours_select, pattern="^hours_\\d+$")],
            HOURS_VALUE: [MessageHandler(filters.TEXT & ~filters.COMMAND, hours_value)],
            HOURS_NOTE: [MessageHandler(filters.TEXT & ~filters.COMMAND, hours_note)],
            CHANGE_PASS_SELECT: [CallbackQueryHandler(change_pass_select, pattern="^chp_\\d+$")],
            CHANGE_PASS_VALUE: [MessageHandler(filters.TEXT & ~filters.COMMAND, change_pass_value)],
        },
        fallbacks=[CommandHandler("cancel", cancel)],
        per_message=False
    )
